import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from io import BytesIO
import markdown
from bs4 import BeautifulSoup, NavigableString

# Check if running in Streamlit context
def _is_streamlit_context():
    """Check if we're running in a Streamlit app context."""
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        return get_script_run_ctx() is not None
    except:
        return False
import re

URL_PATTERN = re.compile(r"(https?://[^\s<>()]+)")

def sanitize_content_for_word(content):
    """Sanitize markdown content for Word export by removing/replacing problematic characters."""
    if not content:
        return content
    
    # Simple approach: just remove all emoji characters using a comprehensive pattern
    # This avoids replacement issues and duplication
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map symbols
        "\U0001F1E0-\U0001F1FF"  # flags (iOS)
        "\U00002500-\U00002BEF"  # chinese char
        "\U00002702-\U000027B0"  # dingbats
        "\U000024C2-\U0001F251"  # enclosed characters
        "\U0001f926-\U0001f937"  # additional emojis
        "\U00010000-\U0010ffff"  # supplementary multilingual plane
        "\u2640-\u2642"         # gender symbols
        "\u2600-\u2B55"         # misc symbols
        "\u200d"                # zero width joiner
        "\u23cf"                # eject symbol
        "\u23e9"                # fast forward
        "\u231a"                # watch
        "\ufe0f"                # variation selector
        "\u3030"                # wavy dash
        "]+", flags=re.UNICODE
    )
    
    # Remove all emojis and replace with space, then clean up
    content = emoji_pattern.sub(' ', content)
    
    # Replace problematic characters that might cause Word issues
    content = content.replace('\u2019', "'")  # Right single quotation mark
    content = content.replace('\u2018', "'")  # Left single quotation mark
    content = content.replace('\u201C', '"')  # Left double quotation mark
    content = content.replace('\u201D', '"')  # Right double quotation mark
    content = content.replace('\u2013', '-')  # En dash
    content = content.replace('\u2014', '--') # Em dash
    content = content.replace('\u2026', '...') # Horizontal ellipsis
    
    # Clean up multiple consecutive spaces and newlines
    content = re.sub(r' {2,}', ' ', content)
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    return content.strip()


def _set_table_borders(table):
    """Add borders to all cells in the table."""
    try:
        tbl = table._tbl
        for row in tbl.tr_lst:
            for cell in row.tc_lst:
                # Get or create table cell properties
                tcPr = cell.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell.append(tcPr)
                
                # Create borders element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Define border style (single line, black, 0.5pt)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # 0.5pt
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    except Exception as e:
        # If border setting fails, continue without borders
        pass


def _linkify_plain_urls(soup: BeautifulSoup) -> None:
    """Wrap bare URLs in anchor tags so they become hyperlinks."""
    for text_node in soup.find_all(string=URL_PATTERN):
        if not isinstance(text_node, NavigableString):
            continue
        parent = text_node.parent
        if parent and parent.name in {'a', 'code', 'pre'}:
            continue
        text_value = str(text_node)
        parts = URL_PATTERN.split(text_value)
        if len(parts) <= 1:
            continue
        new_nodes = []
        for index, part in enumerate(parts):
            if index % 2 == 0:
                if part:
                    new_nodes.append(NavigableString(part))
            else:
                link_text = part
                trailing = ''
                while link_text and link_text[-1] in '.,);':
                    trailing = link_text[-1] + trailing
                    link_text = link_text[:-1]
                if not link_text:
                    new_nodes.append(NavigableString(part))
                    continue
                link = soup.new_tag('a', href=link_text)
                link.string = link_text
                new_nodes.append(link)
                if trailing:
                    new_nodes.append(NavigableString(trailing))
        for node in reversed(new_nodes):
            text_node.insert_after(node)
        text_node.extract()




def _add_table_to_doc(doc, table_element):
    """Add a properly formatted table to the Word document."""
    try:
        # Extract table data
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Get headers and data rows
        headers = []
        data_rows = []

        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            if not cells:
                continue

            cell_texts = [cell.get_text().strip() for cell in cells]

            # First row with th elements or first row is considered header
            if row.find_all('th') or row_idx == 0:
                headers = cell_texts
            else:
                data_rows.append(cell_texts)

        # If no data rows, don't create table
        if not data_rows and not headers:
            return

        # Create Word table
        num_cols = max(len(headers) if headers else 0,
                      max(len(row) for row in data_rows) if data_rows else 0)
        num_rows = (1 if headers else 0) + len(data_rows)

        if num_rows == 0 or num_cols == 0:
            return

        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = 'Table Grid'  # Use Table Grid for better borders
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths to distribute evenly
        for col in word_table.columns:
            col.width = Inches(6.5 / num_cols)  # Distribute across page width

        row_idx = 0

        # Add headers if they exist
        if headers:
            header_row = word_table.rows[row_idx]
            for col_idx, header_text in enumerate(headers[:num_cols]):
                cell = header_row.cells[col_idx]
                cell.text = header_text
                # Format header text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.name = 'Arial'
                
                # Add header cell shading (light gray)
                try:
                    from docx.oxml.ns import nsdecls, parse_xml
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E7E6E6"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                except:
                    pass  # Continue without shading if it fails
                    
            row_idx += 1

        # Add data rows
        for data_row in data_rows:
            if row_idx >= len(word_table.rows):
                break
            word_row = word_table.rows[row_idx]
            for col_idx, cell_text in enumerate(data_row[:num_cols]):
                if col_idx < len(word_row.cells):
                    cell = word_row.cells[col_idx]
                    cell.text = cell_text
                    # Format data cell text
                    for paragraph in cell.paragraphs:
                        # Left align first column (usually IDs), center others for numeric data
                        if col_idx == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            # Check if content looks numeric or date-like for center alignment
                            text = cell_text.strip()
                            if (text.replace('-', '').replace('+', '').replace(' ', '').replace('days', '').replace('Day', '').isdigit() or
                                any(month in text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 
                                                               'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'])):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Arial'
            row_idx += 1

        # Apply borders to the table
        _set_table_borders(word_table)

        # Add some spacing after table
        doc.add_paragraph()

    except Exception as e:
        try:
            if _is_streamlit_context():
                st.error(f"Error adding table to document: {str(e)}")
        except:
            pass

def _is_nested_element(element, processed_elements):
    """Check if an element is nested inside another element we're already processing."""
    parent = element.parent
    while parent:
        if parent in processed_elements:
            return True
        # Check if parent is a type we process at top level
        if parent.name in ['ul', 'ol', 'table', 'blockquote']:
            return True
        parent = parent.parent
    return False


def convert_md_to_docx(md_content):
    """Convert markdown content to a Word document with proper table support."""
    try:
        # Sanitize content for Word compatibility
        sanitized_content = sanitize_content_for_word(md_content)
        
        # Convert markdown to HTML with table extension
        html_content = markdown.markdown(sanitized_content, extensions=['tables', 'nl2br', 'fenced_code'])

        # Create a new Word document
        doc = Document()

        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        _linkify_plain_urls(soup)

        # Process only top-level elements to avoid duplicates
        processed_elements = set()
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'hr', 'blockquote', 'pre', 'code']):
            # Skip elements that are nested inside other elements we're processing
            if _is_nested_element(element, processed_elements):
                continue
                
            processed_elements.add(element)
            
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Add heading
                level = int(element.name[1])
                heading_text = element.get_text().strip()
                if heading_text:
                    doc.add_heading(heading_text, level=level)
            elif element.name == 'p':
                # Skip paragraphs that are inside list items
                if element.parent and element.parent.name == 'li':
                    continue
                # Add paragraph (skip empty paragraphs)
                text = element.get_text().strip()
                if text and text not in ['---', '***', '___']:  # Skip horizontal rules
                    # Check if paragraph contains bold or italic formatting
                    p = doc.add_paragraph()
                    _add_formatted_text(p, element)
            elif element.name == 'ul':
                # Add unordered list
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(paragraph, li)
            elif element.name == 'ol':
                # Add ordered list
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Number')
                    _add_formatted_text(paragraph, li)
            elif element.name == 'table':
                # Add table with proper formatting
                _add_table_to_doc(doc, element)
            elif element.name == 'hr':
                # Add horizontal rule as a line break
                doc.add_paragraph('─' * 50)
            elif element.name == 'blockquote':
                # Add blockquote as indented paragraph
                quote_text = element.get_text().strip()
                if quote_text:
                    p = doc.add_paragraph(quote_text)
                    p.style = 'Quote'
            elif element.name in ['pre', 'code']:
                # Skip code elements that are inside paragraphs or list items
                if element.parent and element.parent.name in ['p', 'li']:
                    continue
                # Add code blocks as monospace text
                code_text = element.get_text().strip()
                if code_text:
                    p = doc.add_paragraph(code_text)
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)

        return doc
    except Exception as e:
        try:
            if _is_streamlit_context():
                st.error(f"Error converting to Word document: {str(e)}")
        except:
            pass
        return None


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink run to a paragraph."""
    if not url:
        paragraph.add_run(text)
        return
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        
        # Add hyperlink style
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        r_pr.append(r_style)
        
        # Add explicit blue color for hyperlinks
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Word's default hyperlink blue
        r_pr.append(color)
        
        # Add underline
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        r_pr.append(underline)
        
        new_run.append(r_pr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        paragraph.add_run(text)


def _add_formatted_text(paragraph, element):
    """Add formatted text (bold, italic, hyperlinks) to a paragraph."""
    try:
        # Directly handle plain text nodes
        if isinstance(element, NavigableString):
            text = str(element)
            if not text:
                return
            parts = URL_PATTERN.split(text)
            if len(parts) <= 1:
                paragraph.add_run(text)
                return
            for index, part in enumerate(parts):
                if index % 2 == 0:
                    if part:
                        paragraph.add_run(part)
                else:
                    link_text = part
                    trailing = ''
                    while link_text and link_text[-1] in '.,);':
                        trailing = link_text[-1] + trailing
                        link_text = link_text[:-1]
                    if not link_text:
                        paragraph.add_run(part)
                        continue
                    _add_hyperlink(paragraph, link_text, link_text)
                    if trailing:
                        paragraph.add_run(trailing)
            return

        if not hasattr(element, 'name'):
            paragraph.add_run(str(element))
            return

        if element.name == 'a' and element.get('href'):
            link_text = element.get_text().strip() or element.get('href')
            _add_hyperlink(paragraph, link_text, element.get('href'))
            return

        for content in element.contents if hasattr(element, 'contents') else []:
            if isinstance(content, NavigableString):
                # Process text nodes for inline URLs
                text = str(content)
                if text:
                    parts = URL_PATTERN.split(text)
                    if len(parts) > 1:
                        for index, part in enumerate(parts):
                            if index % 2 == 0:
                                if part:
                                    paragraph.add_run(part)
                            else:
                                link_text = part
                                trailing = ''
                                while link_text and link_text[-1] in '.,);':
                                    trailing = link_text[-1] + trailing
                                    link_text = link_text[:-1]
                                if not link_text:
                                    paragraph.add_run(part)
                                    continue
                                _add_hyperlink(paragraph, link_text, link_text)
                                if trailing:
                                    paragraph.add_run(trailing)
                    else:
                        paragraph.add_run(text)
            elif hasattr(content, 'name'):
                if content.name in {'strong', 'b'}:
                    # Check if bold element contains links
                    if content.find('a'):
                        _add_formatted_text(paragraph, content)
                    else:
                        run = paragraph.add_run(content.get_text())
                        run.bold = True
                elif content.name in {'em', 'i'}:
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                elif content.name == 'a' and content.get('href'):
                    link_text = content.get_text().strip() or content.get('href')
                    _add_hyperlink(paragraph, link_text, content.get('href'))
                else:
                    _add_formatted_text(paragraph, content)
    except Exception:
        paragraph.add_run(element.get_text() if hasattr(element, 'get_text') else str(element))




def get_docx_bytes(doc):
    """Convert a docx Document object to bytes."""
    try:
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()
    except Exception as e:
        try:
            if _is_streamlit_context():
                st.error(f"Error generating Word document bytes: {str(e)}")
        except:
            pass
        return None
