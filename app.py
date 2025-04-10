import streamlit as st
from docx import Document
from io import BytesIO
import markdown
from bs4 import BeautifulSoup

def set_page_config():
    st.set_page_config(
        page_title="Markdown Viewer",
        page_icon="favicon.ico",
        layout="wide",
        initial_sidebar_state="expanded"
    )

def add_custom_css():
    st.markdown("""
        <style>
        .main {
            background-color: #f5f5f5;
        }
        .stTextArea textarea {
            font-family: 'Courier New', monospace;
            font-size: 16px;
        }
        </style>
    """, unsafe_allow_html=True)

def display_header():
    col1, col2 = st.columns([1, 18])
    with col1:
        st.image("logo.png", width=100)
    with col2:
        st.title("MD Text Viewer & Converter")

def convert_md_to_docx(md_content):
    """Convert markdown content to a Word document."""
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Create a new Word document
    doc = Document()
    
    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
        if element.name in ['h1', 'h2', 'h3']:
            # Add heading
            level = int(element.name[1])
            doc.add_heading(element.get_text().strip(), level=level)
        elif element.name == 'p':
            # Add paragraph
            doc.add_paragraph(element.get_text().strip())
        elif element.name == 'ul':
            # Add unordered list
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text().strip(), style='List Bullet')
        elif element.name == 'ol':
            # Add ordered list
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text().strip(), style='List Number')
    
    return doc

def get_docx_bytes(doc):
    """Convert a docx Document object to bytes."""
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

def main():
    set_page_config()
    add_custom_css()
    display_header()
    
    # Add some instructions
    # st.write("Enter MD text in the text area below. The app will render it as Markdown in real-time.")
    
    with st.expander("Markdown Tips", expanded=False):
        st.markdown("""
        **Markdown Tips:**
        - Use `#` for headers (e.g., `# Header 1`)
        - Use `*` or `_` for *italics* and `**` or `__` for **bold**
        - Use `-` or `*` for bullet lists
        - Use `1.`, `2.`, etc. for numbered lists
        - Use `[text](url)` for links
        - Use `` `code` `` for inline code
        - Add two spaces at the end of a line to create a line break  
        - Multiple spaces will be preserved in the output
        """)
    
    # Create a text area for user input
    user_text = st.text_area("Enter MD text here:", height=200)
    
    # Display the rendered markdown and add download button
    if user_text:
        with st.container(border=True):
            st.subheader("Markdown Output:")
            st.markdown(user_text)
        
        # Add download buttons directly
        # Convert to Word and offer download
        doc = convert_md_to_docx(user_text)
        st.download_button(
            label="📄 Download as Word",
            data=get_docx_bytes(doc),
            file_name="document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.download_button(
            label="📥 Download as Markdown",
            data=user_text,
            file_name="document.md",
            mime="text/markdown"
        )
    
if __name__ == "__main__":
    main() 